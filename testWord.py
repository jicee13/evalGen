from docx import Document
from docx.shared import Pt

reviewClass = 'KG. 2A'
students = ['jayce','barad']

keyPhrases = [
['Physical Development:','Fine Motor Skills:','Gross Motor Skills:'],
['Emotional Development:','Attitude:','Playing With Others:'],
['Cognitive Development:','Problem Solving:','Cognitive Skills:']]


for studentName in students:

    document = Document()

    sections = document.sections
    for section in sections:
        section.top_margin = 1


    heading = document.add_heading()
    headRun = heading.add_run()
    headRun.add_picture('pat.png')
    headRun.add_text('Student Evaluation First Quarter')
    font = headRun.font
    heading.alignment = 1
    font.size = Pt(25)


    paragraph = document.add_paragraph()
    paragraph.add_run('Student\'s Name: '  + studentName)
    paragraph.alignment = 1
    paragraph.add_run('\n')
    paragraph.add_run('Class: ' + reviewClass)
    paragraph.alignment = 1


    for phraseBunch in keyPhrases:
        paragraph = document.add_paragraph()
        topRun = paragraph.add_run(phraseBunch[0])
        topFont = topRun.font
        topFont.bold = True
        topFont.underline = True
        paragraph = document.add_paragraph()
        paragraph.add_run(phraseBunch[1]).italic = True
        paragraph = document.add_paragraph()
        paragraph.add_run(phraseBunch[2]).italic = True

    paragraph = document.add_paragraph()
    paragraph.add_run('\n')
    paragraph.add_run('\n')
    paragraph.add_run('Commented By: T. Jayce').bold = True
    paragraph.alignment = 2

    document.save(studentName + '.docx')
